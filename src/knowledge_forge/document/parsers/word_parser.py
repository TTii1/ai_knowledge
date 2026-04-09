"""Word 文档解析器 - 基于 python-docx，支持中英文标题识别"""

import logging
import re
from pathlib import Path

from knowledge_forge.document.parsers import BaseParser, ParsedDocument, DocumentSection, Table

logger = logging.getLogger(__name__)

# 中文标题样式映射
HEADING_STYLE_MAP = {
    "heading 1": 1, "heading 2": 2, "heading 3": 3,
    "heading 4": 4, "heading 5": 5, "heading 6": 6,
    "标题 1": 1, "标题 2": 2, "标题 3": 3,
    "标题 4": 4, "标题 5": 5, "标题 6": 6,
    "标题1": 1, "标题2": 2, "标题3": 3,
    "标题4": 4, "标题5": 5, "标题6": 6,
    "toc 1": 1, "toc 2": 2, "toc 3": 3,
    "toc1": 1, "toc2": 2, "toc3": 3,
}

# 英文 Heading 样式正则
EN_HEADING_RE = re.compile(r"heading\s*(\d+)", re.IGNORECASE)
CN_HEADING_RE = re.compile(r"标题\s*(\d+)")


def _parse_heading_level(style_name: str) -> int | None:
    """解析标题级别，兼容中英文样式名

    Args:
        style_name: Word 样式名称，如 "Heading 1", "标题 2"

    Returns:
        标题级别 1-6，如果不是标题则返回 None
    """
    if not style_name:
        return None

    # 直接匹配预定义映射
    level = HEADING_STYLE_MAP.get(style_name.lower().strip())
    if level:
        return level

    # 正则匹配
    en_match = EN_HEADING_RE.match(style_name)
    if en_match:
        return int(en_match.group(1))

    cn_match = CN_HEADING_RE.match(style_name)
    if cn_match:
        return int(cn_match.group(1))

    return None


class WordParser(BaseParser):
    """Word (.docx) 文档解析器，支持中英文标题层级和表格提取"""

    supported_extensions = [".docx"]

    async def parse(self, file_path: Path) -> ParsedDocument:
        """解析 Word 文档"""
        from docx import Document

        logger.info("开始解析 Word: %s", file_path)

        doc = Document(str(file_path))
        sections: list[DocumentSection] = []
        tables: list[Table] = []
        full_content_parts: list[str] = []

        current_heading = ""
        current_level = 0
        current_content_parts: list[str] = []

        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""
            heading_level = _parse_heading_level(style_name)

            if heading_level is not None and para.text.strip():
                # 检测到标题 → 保存前一个 section
                if current_content_parts:
                    content = "\n".join(current_content_parts)
                    sections.append(DocumentSection(
                        title=current_heading,
                        content=content,
                        level=current_level,
                    ))
                    full_content_parts.append(content)
                    current_content_parts = []

                current_level = heading_level
                current_heading = para.text.strip()
                current_content_parts.append(para.text)
            else:
                current_content_parts.append(para.text)

        # 保存最后一个 section
        if current_content_parts:
            content = "\n".join(current_content_parts)
            sections.append(DocumentSection(
                title=current_heading,
                content=content,
                level=current_level,
            ))
            full_content_parts.append(content)

        # 提取表格
        for table in doc.tables:
            rows_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                rows_data.append(row_data)
            if rows_data:
                tables.append(Table(
                    headers=rows_data[0],
                    rows=rows_data[1:],
                ))

        full_content = "\n\n".join(full_content_parts)
        title = Path(file_path).stem

        return ParsedDocument(
            title=title,
            content=full_content,
            sections=sections,
            metadata={"has_tables": len(tables) > 0},
            tables=tables,
            source_file=str(file_path),
            total_pages=len(sections),
            file_size=file_path.stat().st_size,
            file_type="docx",
        )
